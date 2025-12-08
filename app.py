from flask import Flask, render_template, request, redirect, url_for, session, send_file, flash
import os
import io, time, logging
import re
import fitz
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests

# ---------------- Load Environment ----------------
load_dotenv()

OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
OPENROUTER_MODEL = os.getenv("OPENROUTER_MODEL", "qwen/qwen-2.5-7b-instruct")
OPENROUTER_BASE_URL = os.getenv("OPENROUTER_BASE_URL", "https://openrouter.ai/api/v1/chat/completions")

if not OPENROUTER_API_KEY:
    raise ValueError("❌ OPENROUTER_API_KEY missing in .env")


app = Flask(__name__)

# --- App Configuration ---
app.config["SECRET_KEY"] = os.getenv("SECRET_KEY", "a-very-secret-key-that-you-should-change")
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///db.sqlite"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

UPLOAD_FOLDER = "uploads"
LOGO_FOLDER = "static/logos"
FONT_FOLDER = "fonts"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(LOGO_FOLDER, exist_ok=True)

db = SQLAlchemy(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = "login"

# --- Utilities ---
logging.basicConfig(level=logging.INFO)



# ---------------- User Model ----------------
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(100), unique=True, nullable=False)
    password = db.Column(db.String(100), nullable=False)
    name = db.Column(db.String(1000))


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))



# ---------------- Extract PDF Text ----------------
def extract_text_from_pdf(path):
    text = ""
    pdf = fitz.open(path)
    for page in pdf:
        text += page.get_text("text")
    return text


# ---------------- FIX BROKEN TAGS ONLY (Preserve Quality) ----------------
def finalize_questions(raw_questions, difficulty):
    cleaned = []
    for q in raw_questions:
        q = q.strip()

        # Fix broken incomplete tags like "(L"
        if re.search(r"\(L\s*$", q):
            q = re.sub(r"\(L.*$", "", q).strip() + f" ({difficulty})"

        cleaned.append(q)
    return cleaned


# ---------------- AI Question Generator ----------------
def generate_ai_questions(text, num_questions=10, difficulty="L1"):
    try:
        prompt = f"""
Generate {num_questions} academic questions from the text below, with a good mix of difficulty levels.

### IMPORTANT: USE BLOOM'S TAXONOMY CORRECTLY
Use ONLY the correct action verbs for the selected level:

L1 (Remember & Understand)
- Define, List, Identify, Name, What is, State, Mention. Assign 2 or 5 marks.

L2 (Apply & Analyze)
- Explain, Describe, Compare, Why, How, Illustrate, Interpret. Assign 5 or 8 marks.

L3 (Evaluate & Create)
- Analyze, Justify, Evaluate, Prove, Derive, Assess, Design. Assign 10 or 15 marks.

### QUESTION DISTRIBUTION:
Generate a mix of questions from L1, L2, and L3.
Aim for a balanced distribution across the levels.

### DO NOT NUMBER THE QUESTIONS
Do NOT add: "1.", "Q1", "a)", "(1)", "-", "*" or any numbering.

### OUTPUT FORMAT RULE
Each question MUST end with the level tag AND the marks tag.
The format MUST be exactly:
<question> (L<level>) [<marks>m]

### EXAMPLES:
What is a neural network? (L1) [5m]
Explain the difference between supervised and unsupervised learning. (L2) [8m]
Design a system to predict stock prices. (L3) [15m]

NO EXCEPTIONS.
Do NOT output incomplete tags like "(L" or "( L".

### TEXT:
{text[:2500]}
"""

        headers = {
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json"
        }

        body = {
            "model": OPENROUTER_MODEL,
            "messages": [
                {"role": "system", "content": "You generate high-quality academic questions using the selected Bloom’s level only."},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.3,
            "max_tokens": 700
        }

        response = requests.post(OPENROUTER_BASE_URL, json=body, headers=headers)

        if response.status_code != 200:
            logging.error(f"OpenRouter API Error: {response.status_code} - {response.text}")
            return ["Error: Could not generate questions due to an API issue."]

        raw = response.json()["choices"][0]["message"]["content"]

        # Split lines & remove blanks
        lines = [line.strip() for line in raw.split("\n") if line.strip()]

        # FIX broken tags ONLY (do NOT override level)
        final = finalize_questions(lines, difficulty)

        return final[:num_questions]

    except Exception as e:
        logging.error(f"AI question generation failed: {e}")
        return ["Error generating AI questions."]


# ---------------- Routes ----------------
@app.route("/")
def index():
    return redirect(url_for("dashboard"))


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form.get("email")
        password = request.form.get("password")
        user = User.query.filter_by(email=email).first()

        if not user or not check_password_hash(user.password, password):
            flash("Please check your login details and try again.")
            return redirect(url_for("login"))

        login_user(user)
        return redirect(url_for("dashboard"))

    return render_template("login.html")


@app.route("/signup", methods=["GET", "POST"])
def signup():
    if request.method == "POST":
        email = request.form.get("email")
        name = request.form.get("name")
        password = request.form.get("password")

        # Check if user already exists
        user = User.query.filter_by(email=email).first()
        if user:
            flash("Email address already exists.")
            return redirect(url_for("signup"))

        # Create new user
        new_user = User(
            email=email,
            name=name,
            password=generate_password_hash(password, method="pbkdf2:sha256")
        )

        # Add to database and log in
        db.session.add(new_user)
        db.session.commit()
        return redirect(url_for("login"))

    return render_template("signup.html")


@app.route("/logout")
@login_required
def logout():
    logout_user()
    return redirect(url_for("login"))


@app.route("/dashboard", methods=["GET", "POST"])
@login_required
def dashboard():
    if request.method == "POST":
        pdf = request.files.get("pdf")
        logo = request.files.get("logo")
        difficulty = request.form.get("difficulty")
        num_questions = int(request.form.get("num_questions"))
        subject = request.form.get("subject")
        college_name = request.form.get("college_name")

        if not pdf:
            flash("Please upload a PDF file.")
            return redirect(request.url)

        filename = secure_filename(pdf.filename)
        path = os.path.join(UPLOAD_FOLDER, filename)
        pdf.save(path)

        logo_path = None
        if logo:
            logo_filename = secure_filename(logo.filename)
            logo_path = os.path.join(LOGO_FOLDER, logo_filename)
            logo.save(logo_path)

        text = extract_text_from_pdf(path)
        questions = generate_ai_questions(text, num_questions, difficulty)

        session["questions"] = questions
        session["subject"] = subject
        session["college_name"] = college_name
        session["logo_path"] = logo_path

        return redirect("/questions")

    return render_template("dashboard.html", name=current_user.name)


@app.route("/questions")
@login_required
def show_questions():
    return render_template("result.html",
                           questions=session.get("questions", []),
                           subject=session.get("subject", "Subject"),
                           college_name=session.get("college_name", "College"))


@app.route("/download_pdf")
@login_required
def download_pdf():
    questions = session.get("questions", [])
    subject = session.get("subject", "Subject")
    college_name = session.get("college_name", "College Name")
    logo_path = session.get("logo_path")

    pdf = FPDF()
    pdf.add_page()

    # Get the absolute path to the project directory
    basedir = os.path.abspath(os.path.dirname(__file__))

    # Add a Unicode-supporting font
    font_regular_path = os.path.join(basedir, FONT_FOLDER, "LiberationSans-Regular.ttf")
    font_bold_path = os.path.join(basedir, FONT_FOLDER, "LiberationSans-Bold.ttf")

    pdf.add_font("Liberation", "", font_regular_path, uni=True)
    pdf.add_font("Liberation", "B", font_bold_path, uni=True)

    if logo_path and os.path.exists(logo_path):
        # Center the logo
        pdf.image(logo_path, x=(pdf.w - 33) / 2, y=8, w=33)
        pdf.ln(30) # Move down to avoid overlap

    pdf.set_font("Liberation", "B", 20)
    pdf.cell(0, 10, college_name, ln=True, align="C")
    pdf.set_font("Liberation", "B", 16)
    pdf.cell(0, 10, f"{subject} - Question Paper", ln=True, align="C")
    pdf.ln(5)

    # Add placeholders for Name and USN
    pdf.set_font("Liberation", "", 12)
    # Using two cells for better alignment
    pdf.cell(0, 10, "Name: " + "_" * 40, border=0, ln=0, align="L")
    pdf.set_x(pdf.w / 2)
    pdf.cell(0, 10, "USN: " + "_" * 40, border=0, ln=1, align="L")
    pdf.ln(10)

    # Draw a line
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(5)

    pdf.set_font("Liberation", "", 12)
    for i, q in enumerate(questions, 1):
        # Ensure we are at the left margin before writing a multi-line cell
        pdf.set_x(pdf.l_margin)
        question_text = f"Q{i}. {q}"
        pdf.multi_cell(0, 8, question_text)
        pdf.ln(2) # Add a small gap after each question

    # Use pdf.output() which returns a bytearray, avoiding manual encoding errors.
    buf = io.BytesIO(pdf.output())
    return send_file(buf, download_name="question_paper.pdf", as_attachment=True)


@app.route("/download_docx")
@login_required
def download_docx():
    subject = session.get("subject", "Subject")
    college_name = session.get("college_name", "College Name")
    logo_path = session.get("logo_path")

    doc = Document()

    # Center the logo
    if logo_path and os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(1.5))

    # Add centered headings
    heading = doc.add_heading(college_name, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_heading = doc.add_heading(f"{subject} - Question Paper", 1)
    sub_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add placeholders for Name and USN
    p_info = doc.add_paragraph()
    p_info.add_run("Name: " + "_" * 50)
    p_info.add_run("\t" * 4) # Add tabs for spacing
    p_info.add_run("USN: " + "_" * 50)

    doc.add_paragraph() # Add a blank line for spacing

    for i, q in enumerate(session.get("questions", []), 1):
        doc.add_paragraph(f"Q{i}. {q}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        download_name="question_paper.docx",
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if __name__ == "__main__":
    with app.app_context():
        db.create_all()
    app.run(debug=True)
